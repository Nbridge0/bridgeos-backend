from pypdf import PdfReader
from docx import Document
from app.image_ai import extract_ocr_from_pdf_pages

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150):
    """
    Splits text into overlapping chunks.
    """

    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start += chunk_size - overlap

    return chunks


def extract_plain_text(file, filename: str) -> str:
    file.seek(0)
    raw = file.read()

    return raw.decode("utf-8", errors="ignore")


def extract_pdf_form_fields(reader) -> str:
    """
    Extracts fillable PDF form fields generically.

    Some invoices/forms store values as AcroForm fields instead of normal page text,
    so page.extract_text() can miss line items, prices, quantities, and totals.
    """

    try:
        fields = reader.get_fields() or {}
    except Exception:
        return ""

    if not fields:
        return ""

    lines = ["PDF form fields:"]

    for field_name, field_data in fields.items():
        value = ""

        if isinstance(field_data, dict):
            value = (
                field_data.get("/V")
                or field_data.get("V")
                or field_data.get("/DV")
                or field_data.get("DV")
                or ""
            )
        else:
            value = field_data

        value = str(value or "").strip()

        if not value:
            continue

        lines.append(f"{field_name}: {value}")

    return "\n".join(lines).strip()


def extract_pdf_text(
    file,
    filename: str = "document.pdf"
) -> str:
    """
    Extracts digital PDF text, form fields, and OCR text.

    OCR is run when:
    - no digital text exists;
    - extracted text is very short; or
    - the document appears financial but contains too few numbers.
    """

    parts = []

    try:
        file.seek(0)
        reader = PdfReader(file)

        for page_index, page in enumerate(
            reader.pages
        ):
            try:
                page_text = str(
                    page.extract_text() or ""
                ).strip()

                if page_text:
                    parts.append(
                        f"PDF page {page_index + 1}:\n"
                        f"{page_text}"
                    )

            except Exception as error:
                print(
                    "PDF PAGE TEXT ERROR:",
                    page_index + 1,
                    type(error).__name__,
                    str(error)
                )

        try:
            form_fields = extract_pdf_form_fields(
                reader
            )

            if form_fields:
                parts.append(form_fields)

        except Exception as error:
            print(
                "PDF FORM FIELD ERROR:",
                type(error).__name__,
                str(error)
            )

    except Exception as error:
        print(
            "PDF TEXT ERROR:",
            type(error).__name__,
            str(error)
        )

    digital_text = "\n\n".join(parts).strip()

    lower_text = digital_text.lower()

    financial_markers = [
        "invoice",
        "receipt",
        "subtotal",
        "total",
        "amount",
        "price",
        "unit price",
        "vat",
        "tax",
        "balance",
        "paid",
        "payment",
        "quantity",
        "qty",
        "currency",
        "usd",
        "eur",
        "gbp",
        "$",
        "€",
        "£"
    ]

    looks_financial = any(
        marker in lower_text
        for marker in financial_markers
    )

    digit_count = sum(
        character.isdigit()
        for character in digital_text
    )

    should_run_ocr = (
        not digital_text
        or len(digital_text) < 300
        or (
            looks_financial
            and digit_count < 10
        )
    )

    if should_run_ocr:
        try:
            file.seek(0)

            ocr_text = extract_ocr_from_pdf_pages(
                file=file,
                filename=filename,
                max_pages=30
            )

            ocr_text = str(
                ocr_text or ""
            ).strip()

            if (
                ocr_text
                and ocr_text != "NO_READABLE_TEXT"
            ):
                parts.append(
                    "Scanned PDF OCR text:\n"
                    + ocr_text
                )

        except Exception as error:
            print(
                "PDF OCR FALLBACK ERROR:",
                type(error).__name__,
                str(error)
            )

    try:
        file.seek(0)
    except Exception:
        pass

    final_text = "\n\n".join(parts).strip()

    print(
        "PDF EXTRACTION FINAL:",
        {
            "filename": filename,
            "characters": len(final_text),
            "digital_characters": len(digital_text),
            "digit_count": sum(
                character.isdigit()
                for character in final_text
            ),
            "ocr_attempted": should_run_ocr,
            "preview": final_text[:500]
        }
    )

    return final_text

def extract_docx_text(file) -> str:
    """
    Extracts paragraphs, tables, headers, footers,
    text boxes, and other Word XML text.

    Repeated amounts are preserved because repeated values
    are meaningful in invoices and financial tables.
    """

    file.seek(0)
    document = Document(file)

    parts = []

    def add_text(value):
        clean = " ".join(
            str(value or "").split()
        ).strip()

        if clean:
            parts.append(clean)

    for paragraph in document.paragraphs:
        add_text(paragraph.text)

    def extract_table(table):
        for row in table.rows:
            row_values = []

            for cell in row.cells:
                cell_values = []

                for paragraph in cell.paragraphs:
                    clean = " ".join(
                        str(
                            paragraph.text or ""
                        ).split()
                    ).strip()

                    if clean:
                        cell_values.append(clean)

                for nested_table in cell.tables:
                    nested_rows = []

                    for nested_row in nested_table.rows:
                        nested_cells = []

                        for nested_cell in nested_row.cells:
                            nested_text = " ".join(
                                str(
                                    nested_cell.text or ""
                                ).split()
                            ).strip()

                            if nested_text:
                                nested_cells.append(
                                    nested_text
                                )

                        if nested_cells:
                            nested_rows.append(
                                " | ".join(nested_cells)
                            )

                    if nested_rows:
                        cell_values.extend(
                            nested_rows
                        )

                cell_text = " ".join(
                    cell_values
                ).strip()

                if cell_text:
                    row_values.append(cell_text)

            if row_values:
                add_text(
                    " | ".join(row_values)
                )

    for table in document.tables:
        extract_table(table)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            add_text(paragraph.text)

        for table in section.header.tables:
            extract_table(table)

        for paragraph in section.footer.paragraphs:
            add_text(paragraph.text)

        for table in section.footer.tables:
            extract_table(table)

    try:
        xml_nodes = document.element.body.xpath(
            ".//w:t"
        )

        xml_text = " ".join(
            str(node.text or "").strip()
            for node in xml_nodes
            if str(node.text or "").strip()
        ).strip()

        if xml_text:
            add_text(
                "Additional Word XML text: "
                + xml_text
            )

    except Exception as error:
        print(
            "DOCX XML TEXT EXTRACTION ERROR:",
            type(error).__name__,
            str(error)
        )

    try:
        file.seek(0)
    except Exception:
        pass

    final_text = "\n".join(parts).strip()

    print(
        "DOCX EXTRACTION FINAL:",
        {
            "characters": len(final_text),
            "preview": final_text[:500]
        }
    )

    return final_text
